from django.shortcuts import render,redirect
from .models import Contribution,Comment
from .forms import ContributionForm,EditContributionForm
from users.models import Faculty,Student,Marketing_Coordinator
from django.views.generic import DetailView
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.paginator import Paginator
from django.core.exceptions import ObjectDoesNotExist
from docx.shared import Inches
from datetime import timedelta
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import os
import docx
import datetime
from zipfile import ZipFile


@login_required
def post_contributions(request):
    if request.method == 'GET':
        student_post = Student.objects.get(id=request.user.id)
        faculty = student_post.faculty
        closure_date = student_post.faculty.new_closure_date 
        if datetime.date.today() > closure_date:
            messages.error(request,f'Clousure date ({closure_date.strftime("%d %B %Y")}) has passed. New contributions cannot be posted 😩 but you can still edit the old posts.')
            return redirect(request.META['HTTP_REFERER'])
        else:
            form = ContributionForm(initial={'author':student_post,'faculty':faculty})
            context = {'form' : form }
    else:
        form = ContributionForm(request.POST,request.FILES)
        files = request.FILES.getlist('images')
        context = {'form' : form }
        if form.is_valid:
            instance = form.save(commit=False)
            instance.images = files[0]
            try:
                instance.images2 = files[1]
            except:
                pass
            try:
                instance.images3 = files[2]
            except:
               pass
            instance.save()
            contribution = Contribution.objects.all().first()
            documentation = docx.Document(contribution.word_docu.path)
            try:
                documentation.add_picture(contribution.images.path, width=Inches(1.25))
            except:
                pass
            try:
                documentation.add_picture(contribution.images2.path, width=Inches(1.25))
            except:
                pass
            try:
                documentation.add_picture(contribution.images3.path, width=Inches(1.25))
            except:
                pass
            docx_paras = []
            contribution.title = documentation.paragraphs[0].text
            for para in documentation.paragraphs[1:]:
                docx_paras.append(para.text)
            docx_para = "\n".join(docx_paras)
            contribution.paragraphs = docx_para
            contribution.save()
            word_docx_directory = os.path.dirname(contribution.word_docu.path)
            os.chdir(word_docx_directory)
            file_name = os.path.basename(contribution.word_docu.path)
            zip_file = ZipFile(f'{file_name}.zip','w')
            zip_file.write(file_name)
            zip_file.close()
            return redirect('home')
    return render(request,'post_contributions.html',context)

@login_required
def edit_contribution(request,pk):
    edit_contribution = Contribution.objects.get(id=pk)
    if request.method == 'POST':
        os.remove(f'{edit_contribution.word_docu.path}.zip')
        form  = EditContributionForm(request.POST,request.FILES,instance=edit_contribution)
        form.save() 
        contribution = Contribution.objects.get(id=pk)
        documentation = docx.Document(contribution.word_docu.path)
        documentation._body.clear_content()
        a = documentation.add_heading(contribution.title)
        a.style = documentation.styles['Heading 1']
        try:
            documentation.add_picture(contribution.images.path, width=Inches(1.25))
        except:
            pass
        try:
            documentation.add_picture(contribution.images2.path, width=Inches(1.25))
        except:
            pass
        try:
            documentation.add_picture(contribution.images3.path, width=Inches(1.25))
        except:
            pass
        documentation.add_paragraph(contribution.paragraphs)
        documentation.save(contribution.word_docu.path)
        word_docx_directory = os.path.dirname(contribution.word_docu.path)
        os.chdir(word_docx_directory)
        file_name = os.path.basename(contribution.word_docu.name)
        zip_file = ZipFile(f'{file_name}.zip','w')
        zip_file.write(file_name)
        zip_file.close()
        messages.success(request,'Contribution updated successfully')
        return redirect('detail-contribution',pk=edit_contribution.id)
    else:
         form  = EditContributionForm(instance=edit_contribution)
         context = {'form' : form, 'title':'Edit Contribution'}
    return render(request,'edit_contributions.html',context)


@login_required
def detail_contribution(request,pk):
    contribution = Contribution.objects.get(id=pk)
    documentation = docx.Document(contribution.word_docu.path)
    docx_paras = []
    title_contrib = documentation.paragraphs[0].text
    for para in documentation.paragraphs[1:]:
        docx_paras.append(para.text)
    try:
        comment = Comment.objects.get(post=contribution)
    except ObjectDoesNotExist:
        comment = None
    context ={
        'title_contrib':title_contrib,
        'docx_paras':docx_paras,
        'object' : contribution,
        'comment' : comment,    
        'title' : 'Detail Contribution'
    }
    return render(request,'detail_contributions.html',context)


@login_required
def home(request):
    if request.user.role == 'MC':
        marketing_coordinator = Marketing_Coordinator.objects.get(id=request.user.id)
        contributions = Contribution.objects.filter(faculty=marketing_coordinator.faculty)
        count = contributions.count() 
    elif request.user.role == 'S':
        contributions = Contribution.objects.filter(author = request.user)
        count = contributions.count() 
    else:
        contributions = Contribution.objects.filter(is_selected=True)
        count = contributions.count()
    paginator = Paginator(contributions, 3)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    context = { 
        "contributions" : page_obj,
        'count': count,
        'title' : "Contributions"
    }
    return render(request,'index.html',context)

@login_required
def select_contribution(request,id):
    contribution = Contribution.objects.get(id=id)
    contribution.is_selected = True
    contribution.save()
    messages.success(request,f"Contribution of {contribution.author.name} is selected for the magazine successfully.")
    return redirect(request.META['HTTP_REFERER'])

@login_required
def deselect_contribution(request,id):
    contribution = Contribution.objects.get(id=id)
    contribution.is_selected = False
    contribution.save()
    messages.success(request,f"Contribution of {contribution.author.name} is deselected.")
    return redirect(request.META['HTTP_REFERER'])

@login_required
def delete_contribution(request,pk):
    contribution = Contribution.objects.get(id=pk)
    contribution.delete()
    messages.success(request,"Contribution deleted successfully")
    return redirect(request.META['HTTP_REFERER'])

@login_required
def write_comment(request,pk):
    if request.method =='POST':
        contribution = Contribution.objects.get(id=pk)
        date = contribution.date_posted + timedelta(days = 14)
        if datetime.datetime.now() > date.replace(tzinfo=None):
            messages.success(request,"Can only comment within 14 days of this contribution posted which is passed now.")
        else:
            comment = request.POST['comment']
            commenter = Marketing_Coordinator.objects.get(id = request.user.id)
            Comment.objects.create(post = contribution,commenter=commenter,comment=comment)
            messages.success(request,"Commented Successfully 🥳. See your comment below.")
    return redirect(request.META['HTTP_REFERER'])


@login_required
def edit_comment(request,pk):
    comment = Comment.objects.get(id=pk)
    if request.method == 'GET':
        context ={
            'object' : comment.post,
            'comment' : comment
        }
    else:
        comment_text = request.POST['comment']
        comment.comment = comment_text
        comment.save()
        messages.success(request,"Comment successfully updated 🥳. See your comment below.")
        return redirect('detail-contribution',pk=comment.post.id)
    return render(request,'edit_comment.html',context)


@login_required
def delete_comment(request,pk):
    comment = Comment.objects.get(id=pk)
    comment.delete()
    messages.success(request,"Comment deleted successfully 🥳.")
    return redirect('detail-contribution',pk=comment.post.id)


@login_required
def selected(request):
    marketing_coordinator = Marketing_Coordinator.objects.get(id=request.user.id)
    contributions = Contribution.objects.filter(faculty=marketing_coordinator.faculty,is_selected=True)
    count = contributions.count() 
    context = { 
        "contributions" : contributions,
        'count': str(count) + ' Selected'
    }
    return render(request,'index.html',context)


@login_required
def not_selected(request):
    marketing_coordinator = Marketing_Coordinator.objects.get(id=request.user.id)
    contributions = Contribution.objects.filter(faculty=marketing_coordinator.faculty,is_selected=False)
    count = contributions.count() 
    context = { 
        "contributions" : contributions,
        'count': str(count) + ' Not Selected'
    }
    return render(request,'index.html',context)


@login_required
def search(request):
    search_faculty = request.GET.get('search')
    faculty = Faculty.objects.get(name=search_faculty)
    contributions = Contribution.objects.filter(faculty=faculty,is_selected=True)
    count = contributions.count() 
    context = { 
        "contributions" : contributions,
        'count': str(count)
    }
    return render(request,'index.html',context)


@login_required
def send_mail(request,pk,html = None):
    try:
        username = 'tharlinhtet.mhs@gmail.com'
        password = 'Tharlin123$'
        from_email = 'University Magazine<username>'
        to_email = []
        to_email.append(Student.objects.get(id=pk).email)
        assert isinstance(to_email,list)
        msg = MIMEMultipart('alternative')
        msg['From'] = 'University Magazine <username>'
        msg['To'] = ",".join(to_email)
        msg['Subject'] = request.POST['subject']
        body = request.POST['body']
        txt_part = MIMEText(body, 'plain')
        msg.attach(txt_part)
        if html != None:
            html_part = MIMEText(html,'html')
            msg.attach(html_part)
            
        msg_str = msg.as_string()
        server = smtplib.SMTP(host='smtp.gmail.com',port=587)
        server.ehlo()
        server.starttls()
        server.login(username,password)
        server.sendmail(from_email,to_email,msg_str)
        server.quit()
        messages.success(request,"Gmail sent successfully 🥳 .")
    except :
        messages.error(request,"Something is wrong. Cannot send mail 😤")
    return redirect(request.META['HTTP_REFERER'])

def overview(request):
    contirbution_count = Contribution.objects.all().count()
    contributor_count = Student.objects.all().count()
    comment_count = Comment.objects.all().count()
    context = {
        'title' : 'Overview',
        'total_contribution' : contirbution_count,
        'contributor' : contributor_count,
        'comment' : comment_count,
        'percent' : 20
    }
    return render(request,'overview.html',context)